import os
import logging
import asyncio
import json
import re
import base64
import io
import pytz
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
import anthropic

# ─────────────────────────────────────────────
# CONFIGURACIÓN
# ─────────────────────────────────────────────
TELEGRAM_TOKEN    = os.environ.get('TELEGRAM_TOKEN_PENAL')
ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY')
TELEGRAM_CHAT_ID  = os.environ.get('TELEGRAM_CHAT_ID')
GMAIL_USER        = os.environ.get('GMAIL_USER')
GOOGLE_TOKEN_B64  = os.environ.get('GOOGLE_TOKEN_B64')
TIMEZONE          = 'Europe/Madrid'
SCOPES            = [
    'https://www.googleapis.com/auth/gmail.send',
]

logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)
logger = logging.getLogger(__name__)

claude_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

# Historial de conversación por sesión
conversation_history = []
MAX_HISTORY = 20

# Estado del escrito en curso
escrito_en_curso = {}

# ─────────────────────────────────────────────
# GMAIL API
# ─────────────────────────────────────────────
def get_gmail_service():
    if not GOOGLE_TOKEN_B64:
        return None
    try:
        token_data = json.loads(base64.b64decode(GOOGLE_TOKEN_B64).decode())
        creds = Credentials.from_authorized_user_info(token_data, SCOPES)
        if creds.expired and creds.refresh_token:
            creds.refresh(Request())
        return build('gmail', 'v1', credentials=creds)
    except Exception as e:
        logger.error(f"Error conectando Gmail: {e}")
        return None

def send_email_with_docx(to_addr, subject, body_text, docx_bytes, filename):
    try:
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        from email.mime.base import MIMEBase
        from email import encoders

        service = get_gmail_service()
        if not service:
            return False

        msg = MIMEMultipart()
        msg['Subject'] = subject
        msg['From']    = GMAIL_USER
        msg['To']      = to_addr
        msg.attach(MIMEText(body_text, 'plain', 'utf-8'))

        part = MIMEBase('application', 'octet-stream')
        part.set_payload(docx_bytes)
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename="{filename}"')
        msg.attach(part)

        raw = base64.urlsafe_b64encode(msg.as_bytes()).decode()
        service.users().messages().send(userId='me', body={'raw': raw}).execute()
        logger.info(f"Email con DOCX enviado a {to_addr}")
        return True
    except Exception as e:
        logger.error(f"Error enviando email: {e}")
        return False

# ─────────────────────────────────────────────
# GENERADOR DE DOCX
# ─────────────────────────────────────────────
def generar_docx(texto):
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Parsear y escribir párrafos
    for linea in texto.split('\n'):
        p = doc.add_paragraph()
        linea = linea.strip()
        if not linea:
            continue

        # Detectar encabezados (todo mayúsculas o numerados I., II., etc.)
        if re.match(r'^(I{1,3}V?|VI{0,3}|IX|X{1,3})\.\s', linea) or \
           re.match(r'^(Primera|Segunda|Tercera|Cuarta|Quinta|Única)\s*[\.\-–]', linea):
            run = p.add_run(linea)
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(6)
        elif linea.isupper() and len(linea) > 3:
            run = p.add_run(linea)
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(6)
        else:
            run = p.add_run(linea)
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(4)

        # Fuente
        run.font.name = 'Times New Roman'

        # Justificado para el cuerpo
        if not (run.bold and run.font.size == Pt(11)):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ─────────────────────────────────────────────
# SISTEMA PROMPT
# ─────────────────────────────────────────────
SYSTEM_PROMPT = """Eres el asistente jurídico penal del despacho AP Estudio Jurídico, especializado en derecho penal español.

Tu función es redactar escritos judiciales penales de alta calidad siguiendo el estilo y estructura del despacho.

ESTILO DEL DESPACHO (basado en escritos reales):
- Encabezado: "AL JUZGADO" o "A LA SALA" o "AL TRIBUNAL" en mayúsculas
- Identificación del procurador/abogado: "Doña/Don [NOMBRE], [cargo], bajo la dirección letrada de Doña MAEVA RUIZ VERGÉS / Adrià Paños Ruiz, colegiado del Il·lustre Col·legi de l'Advocacia de Sabadell, ante el Juzgado comparezco y, como mejor proceda en Derecho, DIGO:"
- Estructura numerada: I., II., III. o Primera, Segunda, Tercera para alegaciones
- Citas legales completas con el texto literal del artículo entre comillas
- Citas jurisprudenciales con ECLI cuando sea posible (ej: ECLI:ES:TS:2013:1007)
- Suplico/Solicitud final en párrafo separado con "SOLICITO que tenga por presentado este escrito..."
- Lugar y fecha al final: "Sabadell, [fecha]"
- Firma: "Adrià Paños Ruiz / Letrado"

LEYES QUE CONOCES:
- Código Penal (CP) — especialmente Título Preliminar, Libro I y II
- Ley de Enjuiciamiento Criminal (LECrim)
- Constitución Española (CE) — arts. 17, 24, 25, 120, 124, 125
- Ley Orgánica del Poder Judicial (LOPJ)
- Código Civil (CC) en lo aplicable al proceso penal

JURISPRUDENCIA: Conoces los criterios del Tribunal Supremo, Tribunal Constitucional y TEDH en materia penal. Cita sentencias reales cuando sea relevante.

FLUJO DE TRABAJO:
1. Cuando el abogado te pida un escrito, identifica qué información tienes y qué falta
2. Si falta información esencial, pregunta SOLO lo imprescindible (máximo 3 preguntas a la vez)
3. Cuando tengas suficiente información, redacta el escrito completo
4. Al final del escrito añade la línea: [ESCRITO_LISTO]

RESPONDE SIEMPRE:
- En español jurídico formal
- Con rigor técnico
- Citando artículos y jurisprudencia real
- Siguiendo el estilo del despacho

Hoy es: {today}
"""

# ─────────────────────────────────────────────
# CLAUDE
# ─────────────────────────────────────────────
def ask_claude(user_msg, doc_context=""):
    global conversation_history
    tz    = pytz.timezone(TIMEZONE)
    today = datetime.now(tz).strftime('%d/%m/%Y')
    system = SYSTEM_PROMPT.replace('{today}', today)

    content = user_msg
    if doc_context:
        content = f"El abogado ha adjuntado el siguiente documento:\n\n{doc_context}\n\nMensaje del abogado: {user_msg}"

    conversation_history.append({"role": "user", "content": content})

    if len(conversation_history) > MAX_HISTORY * 2:
        conversation_history = conversation_history[-(MAX_HISTORY * 2):]

    resp = claude_client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=4000,
        system=system,
        messages=conversation_history
    )
    reply = resp.content[0].text.strip()
    conversation_history.append({"role": "assistant", "content": reply})
    return reply

# ─────────────────────────────────────────────
# HANDLERS
# ─────────────────────────────────────────────
async def cmd_start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if TELEGRAM_CHAT_ID and str(update.effective_chat.id) != str(TELEGRAM_CHAT_ID):
        return
    global conversation_history
    conversation_history = []
    await update.message.reply_text(
        "⚖️ *Bot Junior Penal — AP Estudio Jurídico*\n\n"
        "Puedo redactar cualquier escrito penal:\n"
        "• Querellas\n"
        "• Recursos (reforma, apelación, casación)\n"
        "• Escritos de defensa\n"
        "• Conclusiones provisionales/definitivas\n"
        "• Solicitudes y alegaciones\n"
        "• Cualquier otro escrito procesal penal\n\n"
        "Dígame qué necesita. Puede adjuntar documentos y yo le preguntaré lo que me falte.\n\n"
        "Use /nuevo para empezar un escrito nuevo.",
        parse_mode='Markdown'
    )

async def cmd_nuevo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if TELEGRAM_CHAT_ID and str(update.effective_chat.id) != str(TELEGRAM_CHAT_ID):
        return
    global conversation_history, escrito_en_curso
    conversation_history = []
    escrito_en_curso = {}
    await update.message.reply_text(
        "🔄 Sesión reiniciada. ¿Qué escrito necesita redactar?"
    )

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if TELEGRAM_CHAT_ID and str(update.effective_chat.id) != str(TELEGRAM_CHAT_ID):
        return

    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")

    file = await context.bot.get_file(update.message.document.file_id)
    file_bytes = await file.download_as_bytearray()
    filename   = update.message.document.file_name or ""
    caption    = update.message.caption or "He adjuntado este documento."

    doc_text = ""
    try:
        if filename.endswith('.docx'):
            doc = Document(io.BytesIO(bytes(file_bytes)))
            doc_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        elif filename.endswith('.txt'):
            doc_text = bytes(file_bytes).decode('utf-8', errors='ignore')
        elif filename.endswith('.pdf'):
            doc_text = "[PDF adjunto — por favor indique el contenido relevante en texto]"
        else:
            doc_text = f"[Archivo: {filename}]"
    except Exception as e:
        doc_text = f"[No se pudo leer el archivo: {e}]"

    response = ask_claude(caption, doc_context=doc_text[:8000])
    await _process_response(update, context, response)

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if TELEGRAM_CHAT_ID and str(update.effective_chat.id) != str(TELEGRAM_CHAT_ID):
        return

    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action="typing")
    response = ask_claude(update.message.text)
    await _process_response(update, context, response)

async def _process_response(update, context, response):
    # Detectar si el escrito está listo
    if '[ESCRITO_LISTO]' in response:
        texto_escrito = response.replace('[ESCRITO_LISTO]', '').strip()

        # Generar DOCX
        docx_bytes = generar_docx(texto_escrito)

        # Enviar por Telegram como documento
        tz       = pytz.timezone(TIMEZONE)
        fecha    = datetime.now(tz).strftime('%Y%m%d_%H%M')
        filename = f"Borrador_escrito_penal_{fecha}.docx"

        await update.message.reply_text(
            "✅ *Borrador listo.* Enviando documento...",
            parse_mode='Markdown'
        )
        await context.bot.send_document(
            chat_id=update.effective_chat.id,
            document=io.BytesIO(docx_bytes),
            filename=filename,
            caption="📄 Borrador para revisión. Use /nuevo para empezar otro escrito."
        )

        # Enviar también por email
        if GMAIL_USER:
            tz    = pytz.timezone(TIMEZONE)
            today = datetime.now(tz).strftime('%d/%m/%Y')
            send_email_with_docx(
                GMAIL_USER,
                f"⚖️ Borrador escrito penal — {today}",
                f"Adjunto borrador de escrito penal generado el {today}.\n\nRevise y edite según proceda.\n\nAP Estudio Jurídico",
                docx_bytes,
                filename
            )
    else:
        # Respuesta conversacional — dividir si es muy larga
        MAX_LEN = 4000
        if len(response) <= MAX_LEN:
            await update.message.reply_text(response)
        else:
            # Dividir en chunks
            chunks = [response[i:i+MAX_LEN] for i in range(0, len(response), MAX_LEN)]
            for chunk in chunks:
                await update.message.reply_text(chunk)

# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
def main():
    app = Application.builder().token(TELEGRAM_TOKEN).build()

    app.add_handler(CommandHandler("start", cmd_start))
    app.add_handler(CommandHandler("nuevo", cmd_nuevo))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))

    logger.info("✅ Bot Junior Penal iniciado.")
    app.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == '__main__':
    main()
